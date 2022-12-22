from tkinter import *
from tkinter.filedialog import askopenfilename
from docx import Document
import os

# Creating a window
root = Tk()

# Function to open the Word document
def open_doc():
    global my_text
    filename = askopenfilename(filetypes=[('Word files','*.docx')])
    document = Document(filename)
    fullText = []
    for para in document.paragraphs:
        fullText.append(para.text)
    my_text.insert('10.0', '\n'.join(fullText))
def blooms():
    doc = Document(r"C:\\Users\\Rohithi\\Documents\\questions11.docx")
    prefixes = ('A)','B)','C)','D)','a)','b)','c)','d)','(',')','a.','b.','c.','d.')
    
    for para in doc.paragraphs:
        print(para.text)
        dic={1:['Choose','Define','Describe','Find','Give','How','Label','List','Match','Name','Omit','Recall','Relate','Select','Show','Spell','Tell','What','When','Where','Which','Who','Why','Write','choose','define','describe','find','give','how','label','list','match','name','omit','recall','relate','select','show','spell','tell','what','when','where','which','who','why','Write'],
                 2:['Clarify','Classify','Compare','Contrast','Demonstrate','Explain','Extend','Illustrate','Infer','Interpret','Outline','Relate','Rephrase','Show','Summarize','Translate','clarify','classify','compare','contrast','demonstrate','explain','extend','illustrate','infer','interpret','outline','relate','rephrase','show','summarize','translate'],
                 3:['Apply','Build','Choose','Construct','Develop','Experiment with','Identify','Interview','Makeuseof','Model','Organize','Plan','Select','Solve','Utilize','Use','Using','apply','build','choose','construct','develop','experiment with','identify','interview','make use of','model','organize','plan','select','solve','utilize','use','using'],
                 4:['Analyze','Assume','Categorize','Classify','Compare','Conclusion','Contrast','Discover','Dissect','Distinguish','Divide','Examine','Function','Inference','Inspect','Justify','List','Motive','Relationships','Simplify','Survey','Take part in','Test for','Theme','analyze','assume','categorize','classify','compare','conclusion','contrast','discover','dissect','distinguish','divide','examine','function','inference','inspect','justify','list','motive','relationships','simplify','survey','take part in','test for','theme'],
                 5:['Agree','Appraise','Assess','Award','Choose','Compare','Conclude','Criteria','Criticize','Decide','Deduct','Defend','Determine','Disprove','Estimate','Evaluate','Explain','Importance','Influence','Interpret','Judge','Justify','Mark','Measure','Opinion','Perceive','Prioritize','Prove','Rate','Recommend','Rule on','Select','Support','Value','agree','appraise','assess','award','choose','compare','conclude','criteria','criticize','decide','deduct','defend','determine','disprove','estimate','evaluate','explain','importance','influence','Interpret','judge','justify','mark','measure','opinion','perceive','prioritize','prove','rate','recommend','rule on','select','support','value'],
                 6:['Adapt','Brief','Briefly','Build','Change','Choose','Combine','Compile','Compose','Construct','Create','Delete','Design','Develop','Discuss','Elaborate','Estimate','Formulate','Happen','Imagine','Improve','Invent','Make up','Maximize','Minimize','Modify','Original','Originate','Plan','Predict','Propose','Solution','Solve','Suppose','Test','Theory','adapt','brief','briefly','build','change','choose','combine','compile','compose','construct','create','delete','design','develop','discuss','elaborate','estimate','formulate','happen','imagine','improve','invent','make up','maximize','minimize','modify','original','originate','plan','predict','propose','solution','solve','suppose','test','theory']}
        for key,list_of_words in dic.items():
            any_word_in_string = any(word in para.text for word in list_of_words)
            if any_word_in_string:
                print("The blooms level is:- ",key)

    

# Adding a title
root.title("Read Word Document")

# Creating a text box
my_text = Text(root)
my_text.pack()

# Creating a button
open_button = Button(root, text="Open Word Document", command=open_doc)
open_button.pack()
bloom_button = Button(root, text="Blooms Level", command=blooms)
bloom_button.pack()

# Running the main loop
root.mainloop()